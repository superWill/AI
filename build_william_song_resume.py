from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "resumes/William_Song_iOS_Resume.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="DADCE0", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_format(paragraph, before=0, after=4, line=1.08, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_run(paragraph, text, bold=False, size=10.2, color="202124"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(document, text):
    p = document.add_paragraph()
    p.style = "Heading 1"
    set_paragraph_format(p, before=10, after=5, line=1.05, keep=True)
    add_run(p, text, bold=True, size=12.5, color="0B57D0")
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    set_paragraph_format(p, before=0, after=3, line=1.08)
    add_run(p, text, size=9.5)
    return p


def add_job(document, dates, company, role, bullets):
    p = document.add_paragraph()
    set_paragraph_format(p, before=3, after=2, line=1.05, keep=True)
    add_run(p, company, bold=True, size=10.5, color="202124")
    add_run(p, f"  |  {role}", size=9.7, color="3C4043")
    add_run(p, f"  |  {dates}", size=9.5, color="5F6368")
    for bullet in bullets:
        add_bullet(document, bullet)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Normal"].font.size = Pt(10)
styles["Normal"].font.color.rgb = RGBColor.from_string("202124")
styles["Normal"].paragraph_format.space_after = Pt(4)
styles["Normal"].paragraph_format.line_spacing = 1.08

for name in ("Heading 1", "Heading 2", "Heading 3"):
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    style.font.color.rgb = RGBColor.from_string("0B57D0")

bullet_style = styles["List Bullet"]
bullet_style.font.name = "Arial"
bullet_style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
bullet_style.font.size = Pt(9.5)
bullet_style.paragraph_format.left_indent = Inches(0.22)
bullet_style.paragraph_format.first_line_indent = Inches(-0.12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_format(title, after=1, line=1.0)
add_run(title, "William Song", bold=True, size=22, color="202124")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_format(subtitle, after=5, line=1.0)
add_run(subtitle, "iOS 开发工程师  |  大学本科  |  will-313@outlook.com", size=9.7, color="5F6368")

summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_format(summary, after=8, line=1.08)
add_run(
    summary,
    "长期深耕 iOS 客户端与移动产品开发，经历 IM SDK、社交/内容类项目及商业化产品从迭代到维护的完整周期；关注稳定性、性能、架构可维护性和用户体验。",
    size=9.8,
    color="3C4043",
)

add_heading(doc, "专业技能")
skills = [
    "iOS 客户端开发：Swift / Objective-C、UIKit、Auto Layout、常见组件化与模块化实践。",
    "网络与实时通信：IM SDK 接入与客户端消息链路、会话/通知/状态同步等业务场景。",
    "工程质量：崩溃排查、性能优化、包体与启动体验优化、代码评审与持续迭代维护。",
    "产品协作：与产品、设计、后端协同推进需求评审、技术方案、联调测试和上线发布。",
]
for item in skills:
    add_bullet(doc, item)

add_heading(doc, "工作经历")
add_job(
    doc,
    "2026.03 - 至今",
    "亞洲國際科技有限公司",
    "iOS 开发工程师",
    [
        "负责 iOS 端需求开发、线上问题定位与版本迭代，保障核心业务功能稳定交付。",
        "参与客户端工程维护与体验优化，推动代码结构、可维护性和交付效率持续提升。",
    ],
)
add_job(
    doc,
    "2024.04 - 2025.11",
    "Socrates",
    "Socrates 项目 / iOS 开发",
    [
        "负责 Socrates 项目 iOS 端核心功能开发，覆盖需求评审、技术方案拆解、页面搭建、接口联调、测试修复与版本发布。",
        "参与项目从功能迭代到线上维护的完整流程，围绕关键业务链路处理异常状态、数据展示、页面跳转和用户操作反馈。",
        "优化客户端交互体验与稳定性，跟进线上反馈、崩溃日志和测试问题，推动高频问题快速定位与修复。",
        "配合产品、设计、后端完成多轮需求沟通和联调验收，保证 iOS 端实现与业务目标、接口协议和视觉规范一致。",
    ],
)
add_job(
    doc,
    "2019.10 - 2024.03",
    "百幄（北京）网络科技有限公司 / Beem",
    "Beem 项目 / iOS 开发",
    [
        "长期参与 Beem 项目 iOS 端迭代，负责业务模块开发、基础能力维护与版本发布支持。",
        "配合后端与产品团队完成需求评估、接口协议确认、问题排查和上线后的持续优化。",
        "沉淀可复用组件与开发经验，提升多人协作下的代码一致性和交付质量。",
    ],
)
add_job(
    doc,
    "2014.07 - 2019.09",
    "信诺集团",
    "机票业务 / iOS 开发",
    [
        "参与机票业务 iOS 端功能开发，支持航班查询、机票预订、订单管理、支付结果处理和售后相关业务流程。",
        "围绕航旅业务高频场景处理日期筛选、舱位/价格展示、乘机人信息、订单状态流转等复杂页面与数据交互。",
        "配合后端完成机票接口联调、异常状态处理和版本上线支持，保障查询、下单、支付、出票等核心链路稳定。",
        "积累交易型业务、订单状态管理、复杂表单交互和移动端稳定性维护相关开发经验。",
    ],
)
doc.save(OUT)
print(OUT)
